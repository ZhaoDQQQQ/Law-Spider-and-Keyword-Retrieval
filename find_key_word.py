# -*- coding: utf-8 -*-
import doc
import docx
from docx import Document
import os
import time


# 获取文件夹中特定后缀的文件
# 因为下载下来的文件存在doc docx  pdf多种格式
def get_specific_suffix_file(dir_path, store_list=[], suffix_name=".docx"):
    for root, dirs, files in os.walk(dir_path):
        for file in files:
            if os.path.splitext(file)[1] == suffix_name:
                store_list.append(os.path.join(root, file))
    return store_list


# 续写文件  将信息写入txt文件
def write_txt(txt_path, content):
    with open(txt_path, "a", encoding="utf-8") as f:
        f.write(content)


# 检索 文件中是否存在关键字
def find_key_word(target_files, target_word):
    flag = 0
    for i in range(len(target_files)):
        # 每循环100次  输出一下
        if i % 100 == 0:
            print(i)
        # print(docx_files[i])
        # 获取第i个文件的地址
        docx_file_path = target_files[i]
        try:
            # 打开对应文件
            docx_document = Document(docx_file_path)
        except:
            # 打不开  则输出记录一下
            print(docx_file_path)
            continue
        # 法律文件的标题
        title = ""
        total_length = len(docx_document.paragraphs)
        # 一般  关键字在法律的第一章第一条就会出现
        # 所以只检索前1/3  减少时间复杂度  而且也减少了误检
        if total_length >= 15:
            total_length = int(total_length/3.0)
        j = 0
        # 获取法律文件的标题   用于记录 并数据分析
        while j < total_length:
            temp_text = docx_document.paragraphs[j].text
            temp_text = temp_text.replace(" ", "")
            if "（" in temp_text:
                break
            if "(" in temp_text:
                break
            if temp_text == "":
                j += 1
                continue
            title += temp_text
            j += 1
        # print(title)

        # 进行关键字检索
        while j < total_length:
            temp_text = docx_document.paragraphs[j].text
            temp_text = temp_text.replace(" ", "")
            # 若检测到关键字存在  则保存相关信息
            if target_word in temp_text:
                temp_string = ""
                temp_string = temp_string + docx_file_path.split("\\")[2] + "\n"
                temp_string = temp_string + title + "\n"
                temp_string = temp_string + "******************************************************************\n"
                write_txt(txt_path, temp_string)
                flag += 1
                break
            j += 1
    temp_string = "共计：     " + str(flag)
    write_txt(txt_path, temp_string)


if __name__ == '__main__':
    t = time.time()
    start_path = ".\\saver"
    txt_path = "pathAndName.txt"
    # 检索 saver文件夹中所有docx结尾的文件
    # 注：doc文件需要转换成docx文件才能进行检索   python不支持对doc格式文件进行相关检索操作
    docx_files = get_specific_suffix_file(start_path, [], ".docx")
    # 检索该法律文件是否是根据宪法制定的法律
    key_word = "中华人民共和国宪法"
    # 进行检索并记录
    find_key_word(docx_files, key_word)
    print(time.time() - t)





